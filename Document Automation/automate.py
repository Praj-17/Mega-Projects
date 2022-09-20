## Extracting Text from Excel and creating variables
import pandas as pd 
from docx import Document
import os
import shutil 




data = pd.read_csv('data.csv')
def create_variables(data):
    variables = {}
    for i, j in zip(data['variable'], data['key']):
        variables[i] = j
    return variables
def Parser(file, variables):
    document = Document(file)
    for paragraph in document.paragraphs:
        # replace_with_formating(variables,paragraph)
       for var in variables.keys():  
                         paragraph.text = paragraph.text.replace( f'<<{var}>>',  variables[var])
    table_count  = len(document.tables)
    if table_count >0:
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # replace_with_formating(variables,paragraph)
                        for var in variables.keys():  
                         paragraph.text = paragraph.text.replace( f'<<{var}>>',  variables[var])
    document.save(file) 
def replace_with_formating(variables,paragraph):
    for var in variables.keys(): 
            inline = paragraph.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                    if i-1 < 0:
                        continue
                    if i +1 == len(inline):
                        continue
                    if inline[i-1].text + inline[i].text + inline[i+1].text == f'<<{var}>>': 
                        inline[i-1].text = ''.join(str(map(str,inline[i-1:i+1]))).replace( f'<<{var}>>',  variables[var])
    return 1
                    
def check_if_exist(folder):
    if os.path.exists(folder):
        print(f'{folder} already exists')
        return True
    else:
        return False
def add_annex (annex_name, status, file:list, new_name:list):
    try:     
            if status == False:
                os.mkdir('Master Folder/' + annex_name)
                for f,n in zip(file,new_name):
                    Parser(f, create_variables(data))
                    os.rename(f, n)
                    shutil.move(n,'Master Folder/' + annex_name + '/' )
                print('Annex Added Successfully')
    except Exception as e:
            print("Exception: ", e)
def arrange_files():
    status = check_if_exist('Master Folder')
    if status == False:
        os.mkdir('Master Folder')
        Parser('temp_file_1.docx', create_variables(data))
        print('Master file updated sucessfully')
        os.rename('temp_file_1.docx', 'Master File.docx')
        print('Master file renamed sucessfully')
        shutil.move('Master File.docx','Master Folder/')
        print('Master file moved sucessfully')
        status_a = check_if_exist('Annex A')
        add_annex('Annex A',status_a, ['temp_file_2.docx'], ['Sub-A1.docx'])
        status_b = check_if_exist('Annex B')
        add_annex('Annex B',status_b, ['temp_file_3.docx'], ['Sub-B2.docx'])
        status_c = check_if_exist('Annex C')
        add_annex('Annex C',status_c, ['temp_file_4.docx', 'temp_file_5.docx'], ['Sub-C1.docx', 'Sub-C2.docx'])
def restore_base():
    shutil.rmtree('Master Folder')
    for file in os.listdir('Base'):
        shutil.copy2('Base/'+file, file)
if __name__ == '__main__':
    try:
        arrange_files()
    except Exception as e:
        print("Exception: " , e)
        restore_base()
    
    # restore_base() # roll back to original state
        
        
        
        
